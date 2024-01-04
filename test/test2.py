from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

################### CREATING DOCUMENT ################################
document = Document()

################### ORIENTATION SETINGS ################################
current_section = document.sections[-1]
new_width, new_height = current_section.page_height, current_section.page_width
current_section.orientation = WD_ORIENT.LANDSCAPE
current_section.page_width = new_width
current_section.page_height = new_height


################### HEADER SETINGS ################################
def change_header_paragraph(paragraph, font_family, font_size, aligment):
    for run in paragraph.runs:
        run.font.name = font_family
        run.font.size = Pt(font_size)
    paragraph.alignment = aligment


header_paragraph = current_section.header

h = header_paragraph.add_paragraph(f"ЗАТВЕРДЖЕНО")
change_header_paragraph(h, "Times New Roman", 14, WD_PARAGRAPH_ALIGNMENT.LEFT)
h = header_paragraph.add_paragraph(f"Наказ Міністерства освіти і науки,")
change_header_paragraph(h, "Times New Roman", 14, WD_PARAGRAPH_ALIGNMENT.LEFT)
h = header_paragraph.add_paragraph(f"молоді та спорту України")
change_header_paragraph(h, "Times New Roman", 14, WD_PARAGRAPH_ALIGNMENT.LEFT)
h = header_paragraph.add_paragraph(f"від «29» березня 2012 № 384")
change_header_paragraph(h, "Times New Roman", 14, WD_PARAGRAPH_ALIGNMENT.LEFT)
h = header_paragraph.add_paragraph(f"Форма № Н-5.05")
change_header_paragraph(h, "Times New Roman", 14, WD_PARAGRAPH_ALIGNMENT.LEFT)


header_main = header_paragraph.add_paragraph(
    "Міністерство освіти і науки, молоді та спорту України\nІзмаїльський державний гуманітарний університет")
change_header_paragraph(header_main, "Times New Roman", 14, WD_PARAGRAPH_ALIGNMENT.CENTER)

h = header_paragraph.add_paragraph("Освітній ступінь «sdf»")
change_header_paragraph(h, "Times New Roman", 14, WD_PARAGRAPH_ALIGNMENT.LEFT)
h = header_paragraph.add_paragraph("Спеціальність sdf")
change_header_paragraph(h, "Times New Roman", 14, WD_PARAGRAPH_ALIGNMENT.LEFT)
h = header_paragraph.add_paragraph("Предметна спеціалізація sdf")
change_header_paragraph(h, "Times New Roman", 14, WD_PARAGRAPH_ALIGNMENT.LEFT)
h = header_paragraph.add_paragraph("Семестр sdf")
change_header_paragraph(h, "Times New Roman", 14, WD_PARAGRAPH_ALIGNMENT.LEFT)
h = header_paragraph.add_paragraph("Навчальна дисципліна sdf")
change_header_paragraph(h, "Times New Roman", 14, WD_PARAGRAPH_ALIGNMENT.LEFT)
change_header_paragraph(header_main, "Times New Roman", 14, WD_PARAGRAPH_ALIGNMENT.CENTER)

################### FOOTER SETINGS ################################

footer_paragraphs = current_section.footer
footer_paragraph = footer_paragraphs.add_paragraph(
    "Затверджено на засіданні кафедри sfsdfsd")
change_header_paragraph(footer_paragraph, "Times New Roman", 14, WD_PARAGRAPH_ALIGNMENT.LEFT)
footer_paragraph = footer_paragraphs.add_paragraph(
    "Протокол  № 12 від «12» 12 12 р.")
change_header_paragraph(footer_paragraph, "Times New Roman", 14, WD_PARAGRAPH_ALIGNMENT.LEFT)
footer_paragraph = footer_paragraphs.add_paragraph(
    "в.о. зав. кафедри   _______________________________________   fawgwa")
change_header_paragraph(footer_paragraph, "Times New Roman", 14, WD_PARAGRAPH_ALIGNMENT.LEFT)
footer_paragraph = footer_paragraphs.add_paragraph(
    "Екзаменатор         _______________________________________   sasdgasdg")
change_header_paragraph(footer_paragraph, "Times New Roman", 14, WD_PARAGRAPH_ALIGNMENT.LEFT)

################### GENERATOR ################################

for i in range(1,11):
    p = document.add_paragraph(f'ЕКЗАМЕНАЦІЙНИЙ БІЛЕТ  № {i}')
    change_header_paragraph(p, "Times New Roman", 16, WD_PARAGRAPH_ALIGNMENT.CENTER)
    p.runs[0].bold = True



    document.add_page_break()


document.save('demo.docx')
