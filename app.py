import json
import random

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from flask import Flask, render_template, request, send_file

app = Flask(__name__)


@app.route('/')
def hello():
    return render_template('index.html')


@app.route('/help')
def help_page():
    return render_template('help.html')


@app.route('/postrequest', methods=['POST'])
def post_req():
    data = json.loads(request.data)
    print(data)
    ################### CREATING DOCUMENT ################################
    document = Document()

    # data = {'id': '', 'document': {'nameDocument': 'TEST', 'orientation': 'horizontal',
    #                                'head': 'Міністерство освіти і науки, молоді та спорту України\nІзмаїльський державний гуманітарний університет',
    #                                'osvitniyStypin': '1', 'specialnost': '2', 'specializacia': '3', 'semestr': '4',
    #                                'disciplina': '5', 'numberQuestionBilet': '1',
    #                                'questionType': {'templateQuestionType1': '1'}, 'kafedra': '6',
    #                                'protokol': ['7', '8', '9', '10'], 'persKafedra': '11', 'examinator': '12',
    #                                'numberBilet': '13'}, 'questions': {
    #     'openQuestions': {'numberOpenQuestions': '1', 'questionsField': {'openQuestion1': {'id': 1, 'text': '12'}}},
    #     'testQuestions': {'numberTestQuestions': '', 'questionsField': {}}}}

    ################### ORIENTATION SETINGS ################################
    current_section = document.sections[-1]
    if data["document"]["orientation"] == 'horizontal':
        new_width, new_height = current_section.page_height, current_section.page_width
        current_section.orientation = WD_ORIENT.LANDSCAPE
        current_section.page_width = new_width
        current_section.page_height = new_height

    ################### HEADER SETINGS ################################
    def change_header_paragraph(paragraph, font_family, font_size, aligment):
        for run in paragraph.runs:
            run.font.name = font_family
            run.font.size = Pt(font_size)
        paragraph.alignment = aligment

    header_paragraph = current_section.header
    header_main = header_paragraph.add_paragraph(data["document"]["head"])
    change_header_paragraph(header_main, "Times New Roman", 14, WD_PARAGRAPH_ALIGNMENT.CENTER)

    h = header_paragraph.add_paragraph(f"Освітній ступінь «{data['document']['osvitniyStypin']}»")
    change_header_paragraph(h, "Times New Roman", 14, WD_PARAGRAPH_ALIGNMENT.LEFT)

    h = header_paragraph.add_paragraph(f"Спеціальність {data['document']['specialnost']}")
    change_header_paragraph(h, "Times New Roman", 14, WD_PARAGRAPH_ALIGNMENT.LEFT)

    h = header_paragraph.add_paragraph(f"Предметна спеціалізація {data['document']['specializacia']}")
    change_header_paragraph(h, "Times New Roman", 14, WD_PARAGRAPH_ALIGNMENT.LEFT)

    h = header_paragraph.add_paragraph(f"Семестр {data['document']['semestr']}")
    change_header_paragraph(h, "Times New Roman", 14, WD_PARAGRAPH_ALIGNMENT.LEFT)

    h = header_paragraph.add_paragraph(f"Навчальна дисципліна {data['document']['disciplina']}")
    change_header_paragraph(h, "Times New Roman", 14, WD_PARAGRAPH_ALIGNMENT.LEFT)

    change_header_paragraph(header_main, "Times New Roman", 14, WD_PARAGRAPH_ALIGNMENT.CENTER)

    ################### FOOTER SETINGS ################################

    footer_paragraphs = current_section.footer
    footer_paragraph = footer_paragraphs.add_paragraph(
        f"Затверджено на засіданні кафедри {data['document']['kafedra']}")
    change_header_paragraph(footer_paragraph, "Times New Roman", 14, WD_PARAGRAPH_ALIGNMENT.LEFT)
    footer_paragraph = footer_paragraphs.add_paragraph(
        f"Протокол  № {data['document']['protokol'][0]} від «{data['document']['protokol'][1]}» "
        f"{data['document']['protokol'][2]} {data['document']['protokol'][3]} р.")
    change_header_paragraph(footer_paragraph, "Times New Roman", 14, WD_PARAGRAPH_ALIGNMENT.LEFT)
    footer_paragraph = footer_paragraphs.add_paragraph(
        f"в.о. зав. кафедри   _______________________________________   {data['document']['persKafedra']}")
    change_header_paragraph(footer_paragraph, "Times New Roman", 14, WD_PARAGRAPH_ALIGNMENT.LEFT)
    footer_paragraph = footer_paragraphs.add_paragraph(
        f"Екзаменатор         _______________________________________   {data['document']['examinator']}")
    change_header_paragraph(footer_paragraph, "Times New Roman", 14, WD_PARAGRAPH_ALIGNMENT.LEFT)

    ################### GENERATOR ################################

    for i in range(1, int(data['document']['numberBilet']) + 1):
        p = document.add_paragraph(f'ЕКЗАМЕНАЦІЙНИЙ БІЛЕТ  № {i}')
        change_header_paragraph(p, "Times New Roman", 16, WD_PARAGRAPH_ALIGNMENT.CENTER)
        p.runs[0].bold = True

        for j in range(1, int(data['document']['numberQuestionBilet']) + 1):

            if data['document']['questionType'][f"templateQuestionType{j}"] == "1":
                random_number_question = random.randint(1, int(data['questions']['openQuestions']["numberOpenQuestions"]))
                question = data['questions']['openQuestions']["questionsField"][f"openQuestion{random_number_question}"]["text"]
                p = document.add_paragraph(f'{j}) {question}')
                change_header_paragraph(p, "Times New Roman", 18, WD_PARAGRAPH_ALIGNMENT.LEFT)

            else:
                random_number_question = random.randint(1, int(data['questions']['testQuestions']["numberTestQuestions"]))
                question = data['questions']['testQuestions']["questionsField"][f"testQuestion{random_number_question}"]["text"]
                number_answers = data['questions']['testQuestions']["questionsField"][f"testQuestion{random_number_question}"]["numberAnswerQuestion"]
                answers = data['questions']['testQuestions']["questionsField"][f"testQuestion{random_number_question}"]["answers"]
                p = document.add_paragraph(f'{j}) {question}')
                change_header_paragraph(p, "Times New Roman", 18, WD_PARAGRAPH_ALIGNMENT.LEFT)
                p = document.add_paragraph(f'Відповіді ')
                for k in range(1,int(number_answers) + 1):
                    answer = answers[f"answer{k}"]["text"]
                    p.add_run(f"{k}) {answer}     ")
                    change_header_paragraph(p, "Times New Roman", 18, WD_PARAGRAPH_ALIGNMENT.LEFT)
                    print(answer)




        document.add_page_break()
    global name 
    name = f'{data["document"]["nameDocument"]}.docx'
    document.save(name)

    return "OK"

@app.route('/download_file')
def download_file():
    return send_file(name)

if __name__ == '__main__':

    app.run()
